from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_brief_doc(output_path):
    doc = Document()

    # Title
    title = doc.add_heading('Multimodal Graph RAG: Project Brief', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Elevator Pitch
    doc.add_heading('The Elevator Pitch', level=1)
    pitch = doc.add_paragraph(
        "This is a next-generation AI search engine. Unlike standard search (which just looks for words), "
        "this system uses Knowledge Graphs and Multimodal AI to understand the 'context' of your data. "
        "It can search through Text, Images, and Documents simultaneously to give you the smartest possible answers."
    )

    # Tech Stack
    doc.add_heading('The Tech Stack (The Brain)', level=1)
    techs = [
        ("AI Engine", "Google Gemini 1.5 (Handles the 'reasoning' and 'seeing' images)."),
        ("Graph Database", "Neo4j (Stores relationships like 'Brand X makes Product Y')."),
        ("Vector Database", "Weaviate (Handles high-speed visual and text similarity)."),
        ("Backend", "FastAPI (Python) for high performance."),
        ("Frontend", "Next.js for a premium, interactive user experience.")
    ]
    for category, detail in techs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{category}: ")
        run.bold = True
        p.add_run(detail)

    # Key Features
    doc.add_heading('Key Features (The Wow Factor)', level=1)
    features = [
        ("Visual Search", "Upload a photo of a shoe, and the AI 'looks' at it to find the exact match in the catalog—even if you don't know the name."),
        ("Document Intelligence", "Upload a 200-page PDF, and the system extracts recipes or data instantly to answer your questions."),
        ("Graph-Enhanced Chat", "The AI doesn't just guess; it queries a Knowledge Graph to find real relationships (e.g., matching a watch with the right laptop brand)."),
        ("Resilient Architecture", "If one database is down, the system uses 'AI Vision' fallback to keep the search working.")
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # How it Works
    doc.add_heading('How it Works (The Workflow)', level=1)
    steps = [
        "Ingest: Data is 'fed' into the Graph (Neo4j) and the Vector DB (Weaviate).",
        "Retrieve: When a user searches, the system pulls the best matches from all three modalities (Text/Image/PDF).",
        "Reason: Gemini takes those matches and the user's question to generate a human-like, context-aware response."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Bullet')

    # Example
    doc.add_heading('Real-World Example', level=1)
    example = doc.add_paragraph()
    example.add_run("User: ").bold = True
    example.add_run("Uploads an image of a running shoe and asks: 'Is this good for marathons?'\n")
    example.add_run("System: ").bold = True
    example.add_run(
        "1. Sees the image (Multimodal Vision).\n"
        "2. Finds the 'Sentinel Speed Runners' in the database.\n"
        "3. Pulls 'Breathable mesh' and 'Cushioning' details from the Graph.\n"
        "4. Gemini answers: 'Yes! This VelocityX shoe is designed for marathons because of its responsive cushioning.'"
    )

    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_brief_doc("Project_Brief_Summary.docx")
